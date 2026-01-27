import os
from typing import Optional
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from html4docx import HtmlToDocx
from docx import Document
from datetime import datetime
from io import BytesIO
from .resume_models import ResumeModel
from .template_registry import TemplateRegistry
import logging

# Disable fontTools debug logging
logging.getLogger('fontTools').setLevel(logging.WARNING)
logging.getLogger('fontTools.subset').setLevel(logging.WARNING)
logging.getLogger('fontTools.ttLib').setLevel(logging.WARNING)

logger = logging.getLogger(__name__)

class ResumeGenerator:
    def __init__(self, template_path: Optional[str] = None):
        """
        Initialize Resume Generator

        Args:
            template_path: Path to the directory containing resume templates
        """
        if template_path is None:
            # Default to the templates directory
            template_path = os.path.join(
                os.path.dirname(os.path.abspath(__file__)),
                'templates'
            )

        self.template_path = template_path
        self.env = Environment(loader=FileSystemLoader(template_path))
        self.template_registry = TemplateRegistry(template_path)

    def generate_pdf(self, resume_model: ResumeModel, template_name: str = "resume_template_professional.html") -> bytes:
        """
        Generate PDF from ResumeModel

        Args:
            resume_model: Validated ResumeModel instance
            template_name: Name of the template file to use

        Returns:
            PDF file as bytes
        """
        try:
            # Convert ResumeModel to dictionary for template rendering
            template_data = resume_model.model_dump()

            # Prefer filesystem path for PDF rendering when available
            logo_file_path = template_data.pop('logo_file_path', None)
            if logo_file_path:
                template_data['logo_path'] = logo_file_path
            elif 'logo_path' in template_data and template_data['logo_path']:
                logo_path = template_data['logo_path']
                if logo_path.startswith('/'):
                    backend_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                    candidate_path = os.path.join(backend_root, logo_path.lstrip('/'))
                    if os.path.exists(candidate_path):
                        template_data['logo_path'] = candidate_path

            # Load and render template
            template = self.env.get_template(template_name)
            html_output = template.render(resume=template_data)

            # Generate PDF using WeasyPrint
            pdf_bytes = HTML(
                string=html_output,
                base_url=self.template_path
            ).write_pdf()

            if pdf_bytes is None:
                raise Exception("PDF generation failed - WeasyPrint returned None")

            logger.info(f"Successfully generated PDF for {resume_model.name}")
            return pdf_bytes

        except Exception as e:
            logger.error(f"Error generating PDF: {e}")
            raise Exception(f"Failed to generate resume PDF: {str(e)}")

    def generate_html_preview(self, resume_model: ResumeModel, template_name: str = "resume_template_professional.html") -> str:
        """
        Generate HTML preview of the resume

        Args:
            resume_model: Validated ResumeModel instance
            template_name: Name of the template file to use

        Returns:
            HTML string
        """
        try:
            template_data = resume_model.model_dump()

            # Remove filesystem-only path and keep browser-friendly path if provided
            template_data.pop('logo_file_path', None)

            if 'logo_path' in template_data and template_data['logo_path']:
                logo_path = template_data['logo_path']
                if not logo_path.startswith('/'):
                    template_data['logo_path'] = f"/static/{os.path.basename(logo_path)}"

            template = self.env.get_template(template_name)
            html_output = template.render(resume=template_data)

            logger.info(f"Successfully generated HTML preview for {resume_model.name}")
            return html_output

        except Exception as e:
            logger.error(f"Error generating HTML preview: {e}")
            raise Exception(f"Failed to generate resume HTML: {str(e)}")

    def generate_docx(
        self,
        resume_model: ResumeModel,
        template_name: str = "resume_template_professional.html",
        use_direct_modern: bool = True,
        use_direct_minimal: bool = True,
        use_direct_professional: bool = True
    ) -> bytes:
        """
        Generate DOCX from ResumeModel

        Args:
            resume_model: Validated ResumeModel instance
            template_name: Name of the template file to use
            use_direct_modern: Use direct python-docx for Modern template
            use_direct_minimal: Use direct python-docx for Minimal template
            use_direct_professional: Use direct python-docx for Professional template

        Returns:
            DOCX file as bytes
        """
        try:
            if use_direct_professional and template_name == "resume_template_professional.html":
                return self.generate_docx_professional_direct(resume_model)
            if use_direct_modern and template_name == "resume_template_modern.html":
                return self.generate_docx_modern_direct(resume_model)
            if use_direct_minimal and template_name == "resume_template_minimal.html":
                return self.generate_docx_minimal_direct(resume_model)

            # Convert ResumeModel to dictionary for template rendering
            template_data = resume_model.model_dump()

            # Prefer filesystem path for DOCX rendering when available
            logo_file_path = template_data.pop('logo_file_path', None)
            if logo_file_path:
                template_data['logo_path'] = logo_file_path
            elif 'logo_path' in template_data and template_data['logo_path']:
                logo_path = template_data['logo_path']
                if logo_path.startswith('/'):
                    backend_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                    candidate_path = os.path.join(backend_root, logo_path.lstrip('/'))
                    if os.path.exists(candidate_path):
                        template_data['logo_path'] = candidate_path

            # Use DOCX-specific template with inline styles
            docx_template_name = template_name.replace('.html', '_docx.html')

            # Check if DOCX-specific template exists, otherwise fallback to regular template
            docx_template_path = os.path.join(self.template_path, docx_template_name)
            if not os.path.exists(docx_template_path):
                logger.warning(f"DOCX template {docx_template_name} not found, using {template_name}")
                docx_template_name = template_name

            # Load and render template
            template = self.env.get_template(docx_template_name)
            html_output = template.render(resume=template_data)

            # Generate DOCX using html4docx
            parser = HtmlToDocx()
            apply_table_borders = 'professional' in template_name
            if apply_table_borders:
                parser.table_style = 'Table Grid'

            # Create document
            document = Document()

            # Add metadata
            parser.set_initial_attrs(document)
            metadata = parser.metadata
            metadata.set_metadata(
                author=resume_model.name,
                title=f"Resume - {resume_model.title}",
                created=datetime.now().isoformat()
            )

            # Convert HTML to DOCX
            parser.add_html_to_document(html_output, document)

            # Remove leading empty paragraphs
            self._remove_leading_empty_paragraphs(document)

            # Resize logo if present
            if template_data.get('logo_path'):
                self._resize_logo_in_docx(document)

            # Apply table styling (professional uses grid borders)
            self._apply_table_styling(document, add_borders=apply_table_borders)

            # Fix table column widths
            self._fix_table_widths(document)

            # Apply page break controls
            self._apply_page_break_controls(document)

            # Add footer with company info
            if resume_model.footer:
                self._add_footer_to_docx(document, resume_model.footer)

            # Set default font to Calibri
            self._set_default_font(document, 'Calibri')

            # Set page margins to Normal
            self._set_page_margins(document)

            # Save to bytes
            docx_buffer = BytesIO()
            document.save(docx_buffer)
            docx_bytes = docx_buffer.getvalue()
            docx_buffer.close()

            if docx_bytes is None:
                raise Exception("DOCX generation failed")

            logger.info(f"Successfully generated DOCX for {resume_model.name}")
            return docx_bytes

        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            raise Exception(f"Failed to generate resume DOCX: {str(e)}")

    def generate_docx_professional_direct(self, resume_model: ResumeModel) -> bytes:
        """
        Generate a Professional (Cendien) DOCX directly with python-docx.

        This bypasses HTML->DOCX conversion for better layout control matching
        the Cendien professional resume template.
        """
        try:
            from docx.shared import Pt, Inches, RGBColor, Twips
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            document = Document()

            # Page setup: Letter size with Cendien margins
            section = document.sections[0]
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)

            self._set_default_font(document, 'Calibri')

            # Calculate usable width (8.5 - 0.75 - 0.75 = 7.0 inches)
            usable_width = Inches(7.0)

            colors = {
                'black': RGBColor(0x00, 0x00, 0x00),
                'white': RGBColor(0xFF, 0xFF, 0xFF),
                'blue': RGBColor(0x00, 0x70, 0xC0),
                'purple': RGBColor(0x70, 0x30, 0xA0),  # Standard purple for subsection headings
                'gray': RGBColor(0xA6, 0xA6, 0xA6),
                'dark_gray': RGBColor(0x55, 0x55, 0x55),
                'light_blue': 'DAEEF3',  # Hex for shading
                'text': RGBColor(0x00, 0x00, 0x00)
            }

            def safe_text(value):
                if value is None:
                    return ""
                return str(value)

            def join_values(values, sep=", "):
                parts = []
                for item in values or []:
                    text = safe_text(item)
                    if text:
                        parts.append(text)
                return sep.join(parts)

            def add_text(paragraph, text, size, color, bold=False, italic=False):
                run = paragraph.add_run(safe_text(text))
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(size)
                run.font.color.rgb = color
                run.font.name = 'Calibri'
                return run

            def format_paragraph(paragraph, line_spacing=None, space_before=None, space_after=None, alignment=None):
                fmt = paragraph.paragraph_format
                if line_spacing is not None:
                    fmt.line_spacing = line_spacing
                if space_before is not None:
                    fmt.space_before = Pt(space_before)
                if space_after is not None:
                    fmt.space_after = Pt(space_after)
                if alignment is not None:
                    paragraph.alignment = alignment

            def add_spacer(line_spacing=None, space_before=None, space_after=None, style=None):
                spacer = document.add_paragraph()
                if style:
                    spacer.style = style
                format_paragraph(spacer, line_spacing=line_spacing, space_before=space_before, space_after=space_after)
                return spacer

            def set_cell_shading(cell, fill_hex):
                tcPr = cell._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), fill_hex)
                tcPr.append(shd)

            def set_table_borders(table, border_color='000000', border_size='4'):
                tbl = table._element
                tblPr = tbl.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)
                # Remove existing borders
                old_borders = tblPr.find(qn('w:tblBorders'))
                if old_borders is not None:
                    tblPr.remove(old_borders)
                # Add new borders
                tblBorders = OxmlElement('w:tblBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), border_size)
                    border.set(qn('w:color'), border_color)
                    tblBorders.append(border)
                tblPr.append(tblBorders)

            def remove_table_borders(table):
                tbl = table._element
                tblPr = tbl.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)
                old_borders = tblPr.find(qn('w:tblBorders'))
                if old_borders is not None:
                    tblPr.remove(old_borders)
                tblBorders = OxmlElement('w:tblBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    border.set(qn('w:sz'), '0')
                    border.set(qn('w:color'), 'auto')
                    tblBorders.append(border)
                tblPr.append(tblBorders)

            def set_table_cell_margins(table, top=0, left=108, bottom=0, right=108):
                """Set table-level default cell margins (in twips)"""
                tbl = table._element
                tblPr = tbl.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)
                # Remove existing margins
                old_mar = tblPr.find(qn('w:tblCellMar'))
                if old_mar is not None:
                    tblPr.remove(old_mar)
                # Add new margins
                tblCellMar = OxmlElement('w:tblCellMar')
                for side, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
                    mar = OxmlElement(f'w:{side}')
                    mar.set(qn('w:w'), str(val))
                    mar.set(qn('w:type'), 'dxa')
                    tblCellMar.append(mar)
                tblPr.append(tblCellMar)

            def add_section_title(title_text):
                """Add a section title with black background and white text"""
                table = document.add_table(rows=1, cols=1)
                table.autofit = False
                table.columns[0].width = usable_width
                cell = table.cell(0, 0)
                cell.width = usable_width
                set_cell_shading(cell, '000000')
                set_table_borders(table, '000000', '4')
                # Set compact table cell margins (top/bottom: 30 twips ~= 2pt, left/right: 115 twips ~= 8pt)
                set_table_cell_margins(table, top=30, left=115, bottom=30, right=115)
                para = cell.paragraphs[0]
                format_paragraph(para, line_spacing=1.0, space_before=0, space_after=0)
                add_text(para, title_text, size=11, color=colors['white'], bold=True)
                return table

            # ==================== HEADER TABLE ====================
            logo_path = resume_model.logo_file_path or resume_model.logo_path
            header_table = document.add_table(rows=1, cols=2)
            header_table.autofit = False
            header_table.alignment = WD_TABLE_ALIGNMENT.LEFT

            logo_col_width = Inches(1.1)
            info_col_width = usable_width - logo_col_width
            header_table.columns[0].width = logo_col_width
            header_table.columns[1].width = info_col_width

            logo_cell = header_table.cell(0, 0)
            info_cell = header_table.cell(0, 1)
            logo_cell.width = logo_col_width
            info_cell.width = info_col_width

            remove_table_borders(header_table)
            self._remove_cell_borders(logo_cell)
            self._remove_cell_borders(info_cell)

            # Add logo if available
            if logo_path and os.path.exists(logo_path):
                logo_para = logo_cell.paragraphs[0]
                logo_run = logo_para.add_run()
                logo_run.add_picture(logo_path, width=Inches(0.77), height=Inches(0.77))

            # Add header info (name, title, contact)
            # Name line
            name_para = info_cell.paragraphs[0]
            format_paragraph(name_para, line_spacing=1.0, space_before=0, space_after=0)
            add_text(name_para, resume_model.name, size=18, color=colors['black'], bold=True)

            # Title line
            title_para = info_cell.add_paragraph()
            format_paragraph(title_para, line_spacing=1.0, space_before=0, space_after=0)
            add_text(title_para, resume_model.title, size=14, color=colors['blue'], bold=True)

            # Contact line (labels in gray/normal, values in black/bold)
            contact_para = info_cell.add_paragraph()
            format_paragraph(contact_para, line_spacing=1.0, space_before=0, space_after=0)
            first_item = True
            if resume_model.contact.phone:
                add_text(contact_para, "Phone: ", size=12, color=colors['gray'], bold=False)
                add_text(contact_para, resume_model.contact.phone, size=12, color=colors['black'], bold=True)
                first_item = False
            if resume_model.contact.email:
                if not first_item:
                    add_text(contact_para, " | ", size=12, color=colors['gray'], bold=False)
                add_text(contact_para, "Email: ", size=12, color=colors['gray'], bold=False)
                add_text(contact_para, resume_model.contact.email, size=12, color=colors['black'], bold=True)

            # Spacer after header table
            add_spacer(line_spacing=1.0, space_after=0)

            # ==================== PROFESSIONAL SUMMARY ====================
            add_section_title("Professional Summary")
            summary_para = document.add_paragraph()
            format_paragraph(summary_para, line_spacing=1.0, space_before=12, space_after=0)
            add_text(summary_para, resume_model.summary, size=11, color=colors['text'])
            add_spacer(line_spacing=1.0, space_before=6, space_after=6)

            # ==================== KEY EXPERTISE & SKILLS ====================
            add_section_title("Key Expertise, Skills & Core Qualifications")
            # Spacer between header table and skills table
            add_spacer(space_after=0)

            if resume_model.skills:
                skills_table = document.add_table(rows=0, cols=2)
                skills_table.autofit = False
                category_width = Inches(1.65)
                details_width = usable_width - category_width
                skills_table.columns[0].width = category_width
                skills_table.columns[1].width = details_width
                # Skills table should have borders
                set_table_borders(skills_table, '000000', '4')

                for skill in resume_model.skills:
                    row = skills_table.add_row()
                    cat_cell = row.cells[0]
                    det_cell = row.cells[1]
                    cat_cell.width = category_width
                    det_cell.width = details_width

                    # Category cell with light blue shading
                    set_cell_shading(cat_cell, colors['light_blue'])

                    cat_para = cat_cell.paragraphs[0]
                    format_paragraph(cat_para, space_before=3, space_after=3)
                    add_text(cat_para, safe_text(skill.category), size=11, color=colors['black'], bold=True)

                    det_para = det_cell.paragraphs[0]
                    format_paragraph(det_para, space_before=3, space_after=3)
                    add_text(det_para, safe_text(skill.details), size=11, color=colors['text'])

            # Spacers between skills table and Core Competencies header
            add_spacer(line_spacing=1.0, space_after=0)
            add_spacer(line_spacing=1.0, space_after=0)

            # ==================== CORE COMPETENCIES ====================
            add_section_title("Core Competencies")

            for idx, competency in enumerate(resume_model.core_competencies):
                bullet_para = document.add_paragraph(style='List Bullet')
                format_paragraph(
                    bullet_para,
                    line_spacing=1.0,
                    space_before=12 if idx == 0 else 0,
                    space_after=0
                )
                add_text(bullet_para, safe_text(competency.title), size=11, color=colors['black'], bold=True)
                add_text(bullet_para, f" - {safe_text(competency.description)}", size=11, color=colors['text'])

            # Spacers between Core Competencies and Professional Experience
            add_spacer(line_spacing=1.0, space_after=0)
            add_spacer(line_spacing=1.0, space_after=0)

            # ==================== PROFESSIONAL EXPERIENCE ====================
            add_section_title("Professional Experience")
            # Spacer between header table and first job table
            add_spacer(space_after=0)

            for job in resume_model.experience:
                # Job header table (no borders)
                job_table = document.add_table(rows=2, cols=2)
                job_table.autofit = False
                col1_width = Inches(3.9)
                col2_width = usable_width - col1_width
                job_table.columns[0].width = col1_width
                job_table.columns[1].width = col2_width
                remove_table_borders(job_table)

                # Row 0: Company/Location (spans both columns visually)
                company_cell = job_table.cell(0, 0)
                company_cell.merge(job_table.cell(0, 1))
                company_para = company_cell.paragraphs[0]
                format_paragraph(company_para, line_spacing=1.0, space_before=0, space_after=0)
                company_text = safe_text(job.company)
                if job.location:
                    company_text = f"{company_text} - {safe_text(job.location)}"
                add_text(company_para, company_text, size=11, color=colors['black'], bold=True)

                # Row 1: Role | Dates
                role_cell = job_table.cell(1, 0)
                date_cell = job_table.cell(1, 1)

                role_para = role_cell.paragraphs[0]
                format_paragraph(role_para, line_spacing=1.0, space_before=0, space_after=0)
                add_text(role_para, safe_text(job.role), size=11, color=colors['blue'], bold=True)

                date_para = date_cell.paragraphs[0]
                format_paragraph(date_para, line_spacing=1.0, space_before=0, space_after=0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
                start = job.start_date.strftime('%b %Y') if job.start_date else ''
                end = job.end_date.strftime('%b %Y') if job.end_date else 'Present'
                add_text(date_para, f"{start} - {end}", size=11, color=colors['dark_gray'], italic=True)

                # Job summary
                if job.summary:
                    summary_para = document.add_paragraph()
                    format_paragraph(summary_para, line_spacing=1.0, space_before=6, space_after=6)
                    add_text(summary_para, safe_text(job.summary), size=11, color=colors['text'], italic=True)

                # Notable Projects
                if job.notable_projects:
                    proj_heading = document.add_paragraph()
                    format_paragraph(proj_heading, space_before=6, space_after=6)
                    run = proj_heading.add_run("Notable Projects")
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = colors['purple']
                    for project in job.notable_projects:
                        proj_para = document.add_paragraph(style='List Bullet')
                        format_paragraph(proj_para, line_spacing=1.0, space_after=0)
                        add_text(proj_para, safe_text(project.title), size=11, color=colors['black'], bold=True)
                        add_text(proj_para, f" - {safe_text(project.description)}", size=11, color=colors['text'])

                # Responsibilities
                if job.responsibilities:
                    resp_heading = document.add_paragraph()
                    format_paragraph(resp_heading, space_before=6, space_after=6)
                    run = resp_heading.add_run("Responsibilities")
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = colors['purple']
                    for resp in job.responsibilities:
                        resp_para = document.add_paragraph(style='List Bullet')
                        format_paragraph(resp_para, line_spacing=1.0, space_after=0)
                        add_text(resp_para, resp, size=11, color=colors['text'])

                # Accomplishments
                if job.accomplishments:
                    acc_heading = document.add_paragraph()
                    format_paragraph(acc_heading, space_before=6, space_after=6)
                    run = acc_heading.add_run("Accomplishments")
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = colors['purple']
                    for acc in job.accomplishments:
                        acc_para = document.add_paragraph(style='List Bullet')
                        format_paragraph(acc_para, line_spacing=1.0, space_after=0)
                        add_text(acc_para, acc, size=11, color=colors['text'])

                # Environment
                if job.environment:
                    env_para = document.add_paragraph()
                    format_paragraph(env_para, line_spacing=1.0, space_before=6, space_after=12)
                    add_text(env_para, "Environment: ", size=11, color=colors['black'], bold=True)
                    add_text(env_para, join_values(job.environment), size=11, color=colors['text'])

            if resume_model.education:
                add_spacer(space_before=6, space_after=8)
                add_spacer(line_spacing=1.0, space_after=0)

            # ==================== EDUCATION ====================
            add_section_title("Education")
            add_spacer(space_after=0)

            for edu in resume_model.education:
                edu_table = document.add_table(rows=2, cols=2)
                edu_table.autofit = False
                col1_width = Inches(3.9)
                col2_width = usable_width - col1_width
                edu_table.columns[0].width = col1_width
                edu_table.columns[1].width = col2_width
                remove_table_borders(edu_table)

                # Row 0: Institution/Location
                inst_cell = edu_table.cell(0, 0)
                inst_cell.merge(edu_table.cell(0, 1))
                inst_para = inst_cell.paragraphs[0]
                format_paragraph(inst_para, line_spacing=1.0, space_before=0, space_after=0)
                inst_text = safe_text(edu.institution)
                if edu.location:
                    inst_text = f"{inst_text} - {safe_text(edu.location)}"
                add_text(inst_para, inst_text, size=11, color=colors['black'], bold=True)

                # Row 1: Degree | Graduation Date
                degree_cell = edu_table.cell(1, 0)
                grad_cell = edu_table.cell(1, 1)

                degree_para = degree_cell.paragraphs[0]
                format_paragraph(degree_para, line_spacing=1.0, space_before=0, space_after=0)
                add_text(degree_para, safe_text(edu.degree), size=11, color=colors['blue'], bold=True)

                grad_para = grad_cell.paragraphs[0]
                format_paragraph(grad_para, line_spacing=1.0, space_before=0, space_after=0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
                grad_date = edu.graduation_date.strftime('%b %Y') if edu.graduation_date else ''
                if grad_date:
                    add_text(grad_para, f"Graduated: {grad_date}", size=11, color=colors['dark_gray'], italic=True)

                # Relevance
                if edu.relevance:
                    add_spacer(line_spacing=1.0, space_after=0)
                    rel_para = document.add_paragraph()
                    format_paragraph(rel_para, line_spacing=1.0, space_before=0, space_after=0)
                    add_text(rel_para, safe_text(edu.relevance), size=11, color=colors['text'])

            if resume_model.certifications:
                # Spacer between Education and Certifications (avoid oversized gaps)
                add_spacer(space_before=6, space_after=8)

            # ==================== CERTIFICATIONS ====================
            if resume_model.certifications:
                add_section_title("Certifications")
                add_spacer(space_after=0)

                for cert_index, cert in enumerate(resume_model.certifications):
                    cert_table = document.add_table(rows=1, cols=2)
                    cert_table.autofit = False
                    col1_width = Inches(4.5)
                    col2_width = usable_width - col1_width
                    cert_table.columns[0].width = col1_width
                    cert_table.columns[1].width = col2_width
                    remove_table_borders(cert_table)

                    name_cell = cert_table.cell(0, 0)
                    date_cell = cert_table.cell(0, 1)

                    name_para = name_cell.paragraphs[0]
                    format_paragraph(name_para, line_spacing=1.0, space_before=0, space_after=0)
                    add_text(name_para, safe_text(cert.name), size=11, color=colors['black'], bold=True)

                    date_para = date_cell.paragraphs[0]
                    format_paragraph(date_para, line_spacing=1.0, space_before=0, space_after=0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
                    if cert.issue_date:
                        add_text(date_para, cert.issue_date.strftime('%b %Y'), size=11, color=colors['dark_gray'], italic=True)

                    if cert.description:
                        desc_para = document.add_paragraph()
                        format_paragraph(desc_para, line_spacing=1.0, space_after=0)
                        add_text(desc_para, safe_text(cert.description), size=11, color=colors['text'])
                    if cert_index < len(resume_model.certifications) - 1:
                        add_spacer(line_spacing=1.0, space_after=0)

            # ==================== FOOTER ====================
            if resume_model.footer:
                footer = section.footer
                footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                footer_para.clear()
                footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Add top border to footer
                pPr = footer_para._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                top = OxmlElement('w:top')
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), '4')
                top.set(qn('w:space'), '1')
                top.set(qn('w:color'), '000000')
                pBdr.append(top)
                pPr.append(pBdr)

                # Format footer text with bullet separators
                footer_text = resume_model.footer.replace(' | ', '   •   ')
                run = footer_para.add_run(footer_text)
                run.font.name = 'Calibri'
                run.font.size = Pt(8)

            # Save to bytes
            docx_buffer = BytesIO()
            document.save(docx_buffer)
            docx_bytes = docx_buffer.getvalue()
            docx_buffer.close()

            if docx_bytes is None:
                raise Exception("DOCX generation failed")

            logger.info(f"Successfully generated professional DOCX for {resume_model.name}")
            return docx_bytes

        except Exception as e:
            logger.error(f"Error generating professional DOCX: {e}")
            raise Exception(f"Failed to generate professional resume DOCX: {str(e)}")

    def generate_docx_modern_direct(self, resume_model: ResumeModel) -> bytes:
        """
        Generate a Modern DOCX directly with python-docx (experimental).

        This bypasses HTML->DOCX conversion for better layout control.
        """
        try:
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            document = Document()

            # Page setup: legal with Cendien modern margins
            section = document.sections[0]
            section.page_width = Inches(8.5)
            section.page_height = Inches(14)
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)

            # Default to an Office font (Calibri)
            self._set_default_font(document, 'Calibri')

            usable_width = Inches(7.0)

            colors = {
                'black': RGBColor(0x00, 0x00, 0x00),
                'blue': RGBColor(0x00, 0x70, 0xC0),
                'purple': RGBColor(0x70, 0x30, 0xA0),
                'gray': RGBColor(0x80, 0x80, 0x80),
                'dark_gray': RGBColor(0x55, 0x55, 0x55),
                'text': RGBColor(0x26, 0x26, 0x26),
                'light_blue': 'EDF7F9',
                'light_gray': 'F8F9FA'
            }

            def safe_text(value):
                if value is None:
                    return ""
                return str(value)

            def join_values(values, sep=", "):
                parts = []
                for item in values or []:
                    text = safe_text(item)
                    if text:
                        parts.append(text)
                return sep.join(parts)

            def add_text(paragraph, text, size, color, bold=False, italic=False, align=None):
                run = paragraph.add_run(safe_text(text))
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(size)
                run.font.color.rgb = color
                run.font.name = 'Calibri'
                if align is not None:
                    paragraph.alignment = align
                return run

            def format_paragraph(paragraph, line_spacing=None, space_before=None, space_after=None, alignment=None):
                fmt = paragraph.paragraph_format
                if line_spacing is not None:
                    fmt.line_spacing = line_spacing
                if space_before is not None:
                    fmt.space_before = Pt(space_before)
                if space_after is not None:
                    fmt.space_after = Pt(space_after)
                if alignment is not None:
                    paragraph.alignment = alignment

            def set_cell_shading(cell, fill_hex):
                tcPr = cell._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), fill_hex)
                tcPr.append(shd)

            def set_table_borders(table, border_color, border_size):
                tbl = table._element
                tblPr = tbl.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)
                old_borders = tblPr.find(qn('w:tblBorders'))
                if old_borders is not None:
                    tblPr.remove(old_borders)
                tblBorders = OxmlElement('w:tblBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), str(border_size))
                    border.set(qn('w:color'), border_color)
                    border.set(qn('w:space'), '0')
                    tblBorders.append(border)
                tblPr.append(tblBorders)

            def remove_table_borders(table):
                tbl = table._element
                tblPr = tbl.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)
                old_borders = tblPr.find(qn('w:tblBorders'))
                if old_borders is not None:
                    tblPr.remove(old_borders)
                tblBorders = OxmlElement('w:tblBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    border.set(qn('w:sz'), '0')
                    border.set(qn('w:color'), 'auto')
                    tblBorders.append(border)
                tblPr.append(tblBorders)

            def set_cell_borders(cell, border_color, border_size):
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is None:
                    tcBorders = OxmlElement('w:tcBorders')
                    tcPr.append(tcBorders)
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = tcBorders.find(qn(f'w:{border_name}'))
                    if border is None:
                        border = OxmlElement(f'w:{border_name}')
                        tcBorders.append(border)
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), str(border_size))
                    border.set(qn('w:color'), border_color)
                    border.set(qn('w:space'), '0')

            def add_section_title(text):
                table = document.add_table(rows=1, cols=1)
                table.autofit = False
                table.columns[0].width = usable_width
                set_table_borders(table, '0070C0', '4')
                cell = table.cell(0, 0)
                set_cell_shading(cell, '0070C0')
                para = cell.paragraphs[0]
                format_paragraph(para, space_before=1, space_after=1)
                add_text(para, text, size=12, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
                return table

            # Header with optional logo
            logo_path = resume_model.logo_file_path or resume_model.logo_path
            header_table = document.add_table(rows=1, cols=2)
            header_table.autofit = False
            remove_table_borders(header_table)
            self._set_table_column_widths(header_table, [1300, 8798])
            header_table.columns[0].width = Inches(0.9)
            header_table.columns[1].width = usable_width - Inches(0.9)
            logo_cell = header_table.cell(0, 0)
            info_cell = header_table.cell(0, 1)
            self._remove_cell_borders(logo_cell)
            self._remove_cell_borders(info_cell)

            if logo_path and os.path.exists(logo_path):
                logo_para = logo_cell.paragraphs[0]
                logo_run = logo_para.add_run()
                logo_run.add_picture(logo_path, width=Inches(0.7), height=Inches(0.7))

            name_para = info_cell.paragraphs[0]
            format_paragraph(name_para, space_before=6, space_after=2)
            add_text(name_para, resume_model.name, size=18, color=colors['black'], bold=True)

            title_para = info_cell.add_paragraph()
            format_paragraph(title_para, space_before=2, space_after=2)
            add_text(title_para, resume_model.title, size=14, color=colors['blue'], bold=True)

            contact_para = info_cell.add_paragraph()
            format_paragraph(contact_para, space_after=6)
            first_item = True
            if resume_model.contact.phone:
                add_text(contact_para, "Phone: ", size=12, color=colors['gray'], bold=False)
                add_text(contact_para, resume_model.contact.phone, size=12, color=colors['black'], bold=True)
                first_item = False
            if resume_model.contact.email:
                if not first_item:
                    add_text(contact_para, " | ", size=12, color=colors['gray'])
                add_text(contact_para, "Email: ", size=12, color=colors['gray'], bold=False)
                add_text(contact_para, resume_model.contact.email, size=12, color=colors['black'], bold=True)
                first_item = False
            if resume_model.contact.linkedin:
                if not first_item:
                    add_text(contact_para, " | ", size=12, color=colors['gray'])
                add_text(contact_para, "LinkedIn: ", size=12, color=colors['gray'], bold=False)
                add_text(contact_para, resume_model.contact.linkedin, size=12, color=colors['black'], bold=True)
                first_item = False
            if resume_model.contact.github:
                if not first_item:
                    add_text(contact_para, " | ", size=12, color=colors['gray'])
                add_text(contact_para, "GitHub: ", size=12, color=colors['gray'], bold=False)
                add_text(contact_para, resume_model.contact.github, size=12, color=colors['black'], bold=True)

            # Professional Summary
            add_section_title("Professional Summary")
            summary_lines = [line.strip() for line in safe_text(resume_model.summary).split('\\n') if line.strip()]
            if not summary_lines:
                summary_lines = [safe_text(resume_model.summary)]
            for line in summary_lines:
                summary_paragraph = document.add_paragraph()
                format_paragraph(summary_paragraph, line_spacing=1.0, space_before=12, space_after=6)
                add_text(summary_paragraph, line, size=11, color=colors['text'])

            # Key Expertise Areas & Skills
            add_section_title("Key Expertise Areas & Skills")
            if resume_model.skills:
                # Spacer to prevent section title table from merging with the skills table
                document.add_paragraph()
                skills_table = document.add_table(rows=0, cols=2)
                skills_table.autofit = False
                self._set_table_column_widths(skills_table, [2875, 7200])
                skills_table.columns[0].width = Inches(2.0)
                skills_table.columns[1].width = Inches(5.0)
                set_table_borders(skills_table, '31849B', '4')
                for skill in resume_model.skills:
                    row = skills_table.add_row()
                    left = row.cells[0]
                    right = row.cells[1]
                    set_cell_shading(left, colors['light_blue'])
                    left_paragraph = left.paragraphs[0]
                    format_paragraph(left_paragraph, space_before=3, space_after=3)
                    add_text(left_paragraph, safe_text(skill.category), size=11, color=RGBColor(0x00, 0x20, 0x60), bold=True)
                    right_paragraph = right.paragraphs[0]
                    format_paragraph(right_paragraph, space_before=3, space_after=3)
                    add_text(right_paragraph, safe_text(skill.details), size=11, color=colors['text'])
                # Spacer to prevent skills table from merging with the next section title
                document.add_paragraph()

            # Core Competencies
            add_section_title("Core Competencies")
            if resume_model.core_competencies:
                # Spacer to prevent section title table from merging with the competencies table
                document.add_paragraph()
                comp_table = document.add_table(rows=0, cols=2)
                comp_table.autofit = False
                self._set_table_column_widths(comp_table, [5040, 5040])
                comp_table.columns[0].width = Inches(3.5)
                comp_table.columns[1].width = Inches(3.5)
                set_table_borders(comp_table, '3498DB', '8')
                for idx, competency in enumerate(resume_model.core_competencies):
                    if idx % 2 == 0:
                        row = comp_table.add_row()
                    cell = row.cells[idx % 2]
                    set_cell_shading(cell, colors['light_gray'])
                    title_paragraph = cell.paragraphs[0]
                    format_paragraph(title_paragraph, line_spacing=1.0, space_before=6, space_after=4)
                    add_text(title_paragraph, safe_text(competency.title), size=11, color=colors['blue'], bold=True)
                    desc_paragraph = cell.add_paragraph()
                    format_paragraph(desc_paragraph, line_spacing=1.0, space_before=2, space_after=6)
                    add_text(desc_paragraph, safe_text(competency.description), size=11, color=colors['text'])
                # Spacer to prevent competencies table from merging with the next section title
                document.add_paragraph()

            # Professional Experience
            add_section_title("Professional Experience")
            for job in resume_model.experience:
                job_table = document.add_table(rows=2, cols=2)
                job_table.autofit = False
                self._set_table_column_widths(job_table, [5616, 4456])
                job_table.columns[0].width = Inches(3.9)
                job_table.columns[1].width = usable_width - Inches(3.9)
                set_table_borders(job_table, '000000', '4')
                for row in job_table.rows:
                    for cell in row.cells:
                        set_cell_borders(cell, 'FFFFFF', '6')

                company_cell = job_table.cell(0, 0)
                company_cell.merge(job_table.cell(0, 1))
                company_para = company_cell.paragraphs[0]
                company_text = safe_text(job.company)
                if job.location:
                    company_text = f"{company_text} - {safe_text(job.location)}"
                add_text(company_para, company_text, size=12, color=colors['black'], bold=True)

                role_para = job_table.cell(1, 0).paragraphs[0]
                add_text(role_para, safe_text(job.role), size=12, color=colors['blue'], bold=True)

                date_para = job_table.cell(1, 1).paragraphs[0]
                format_paragraph(date_para, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
                start = job.start_date.strftime('%b %Y') if job.start_date else ''
                end = job.end_date.strftime('%b %Y') if job.end_date else 'Present'
                add_text(date_para, f"{start} - {end}", size=11, color=colors['dark_gray'])

                if job.summary:
                    summary_paragraph = document.add_paragraph()
                    format_paragraph(summary_paragraph, line_spacing=1.0, space_before=6, space_after=6)
                    add_text(summary_paragraph, safe_text(job.summary), size=11, color=colors['text'])

                def add_subsection(title, items):
                    if not items:
                        return
                    heading = document.add_paragraph()
                    format_paragraph(heading, line_spacing=1.0, space_before=6, space_after=6)
                    run = heading.add_run(title)
                    run.bold = True
                    run.font.color.rgb = colors['purple']
                    for item in items:
                        bullet = document.add_paragraph(style='List Bullet')
                        add_text(bullet, item, size=11, color=colors['text'])

                if job.notable_projects:
                    add_subsection(
                        "Notable Projects",
                        [f"{safe_text(p.title)}: {safe_text(p.description)}" for p in job.notable_projects]
                    )
                add_subsection("Responsibilities", job.responsibilities or [])
                add_subsection("Accomplishments", job.accomplishments or [])

                if job.environment:
                    env_paragraph = document.add_paragraph()
                    add_text(env_paragraph, "Environment: ", size=11, color=colors['purple'], bold=True)
                    add_text(env_paragraph, join_values(job.environment), size=11, color=colors['text'])

            # Education
            add_section_title("Education")
            for edu in resume_model.education:
                edu_table = document.add_table(rows=2, cols=2)
                edu_table.autofit = False
                self._set_table_column_widths(edu_table, [5616, 4464])
                edu_table.columns[0].width = Inches(3.9)
                edu_table.columns[1].width = usable_width - Inches(3.9)
                set_table_borders(edu_table, '000000', '4')
                for row in edu_table.rows:
                    for cell in row.cells:
                        set_cell_borders(cell, 'FFFFFF', '6')

                inst_cell = edu_table.cell(0, 0)
                inst_cell.merge(edu_table.cell(0, 1))
                inst_para = inst_cell.paragraphs[0]
                inst_text = safe_text(edu.institution)
                if edu.location:
                    inst_text = f"{inst_text} - {safe_text(edu.location)}"
                add_text(inst_para, inst_text, size=12, color=colors['black'], bold=True)

                degree_para = edu_table.cell(1, 0).paragraphs[0]
                add_text(degree_para, safe_text(edu.degree), size=12, color=colors['blue'], bold=True)

                grad_para = edu_table.cell(1, 1).paragraphs[0]
                format_paragraph(grad_para, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
                grad_date = edu.graduation_date.strftime('%b %Y') if edu.graduation_date else ''
                add_text(grad_para, f"Graduated: {grad_date}", size=11, color=colors['dark_gray'])

                if edu.relevance:
                    rel_para = document.add_paragraph()
                    format_paragraph(rel_para, line_spacing=1.0, space_before=6, space_after=6)
                    add_text(rel_para, safe_text(edu.relevance), size=11, color=colors['text'])

            # Certifications
            if resume_model.certifications:
                add_section_title("Certifications")
                for cert in resume_model.certifications:
                    name_paragraph = document.add_paragraph()
                    format_paragraph(name_paragraph, line_spacing=1.0, space_before=6, space_after=6)
                    add_text(name_paragraph, safe_text(cert.name), size=12, color=colors['purple'], bold=True)
                    if cert.description:
                        desc_paragraph = document.add_paragraph()
                        format_paragraph(desc_paragraph, line_spacing=1.0, space_before=6, space_after=6)
                        add_text(desc_paragraph, safe_text(cert.description), size=11, color=colors['text'])

            # Save to bytes
            docx_buffer = BytesIO()
            document.save(docx_buffer)
            docx_bytes = docx_buffer.getvalue()
            docx_buffer.close()

            if docx_bytes is None:
                raise Exception("DOCX generation failed")

            logger.info(f"Successfully generated modern DOCX for {resume_model.name}")
            return docx_bytes

        except Exception as e:
            logger.error(f"Error generating modern DOCX: {e}")
            raise Exception(f"Failed to generate modern resume DOCX: {str(e)}")

    def generate_docx_minimal_direct(self, resume_model: ResumeModel) -> bytes:
        """
        Generate a Minimal DOCX directly with python-docx (experimental).

        This bypasses HTML->DOCX conversion for better layout control.
        """
        try:
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            document = Document()

            # Page setup: letter with 1" margins to match Minimal HTML
            section = document.sections[0]
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

            self._set_default_font(document, 'Calibri')

            colors = {
                'text': RGBColor(0x1A, 0x1A, 0x1A),
                'muted': RGBColor(0x66, 0x66, 0x66),
                'light': RGBColor(0x99, 0x99, 0x99),
                'border': RGBColor(0xE0, 0xE0, 0xE0),
                'body': RGBColor(0x33, 0x33, 0x33),
                'desc': RGBColor(0x55, 0x55, 0x55)
            }

            def safe_text(value):
                if value is None:
                    return ""
                return str(value)

            def join_values(values, sep=", "):
                parts = []
                for item in values or []:
                    text = safe_text(item)
                    if text:
                        parts.append(text)
                return sep.join(parts)

            def add_text(paragraph, text, size, color, bold=False, italic=False, align=None):
                run = paragraph.add_run(safe_text(text))
                run.bold = bold
                run.italic = italic
                run.font.size = Pt(size)
                run.font.color.rgb = color
                run.font.name = 'Calibri'
                if align is not None:
                    paragraph.alignment = align
                return run

            def format_paragraph(paragraph, line_spacing=None, space_before=None, space_after=None, alignment=None):
                fmt = paragraph.paragraph_format
                if line_spacing is not None:
                    fmt.line_spacing = line_spacing
                if space_before is not None:
                    fmt.space_before = Pt(space_before)
                if space_after is not None:
                    fmt.space_after = Pt(space_after)
                if alignment is not None:
                    paragraph.alignment = alignment

            def set_paragraph_bottom_border(paragraph, color_hex, size=4):
                pPr = paragraph._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), str(size))
                bottom.set(qn('w:color'), color_hex)
                pBdr.append(bottom)
                pPr.append(pBdr)

            def set_paragraph_left_border(paragraph, color_hex, size=6):
                pPr = paragraph._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'single')
                left.set(qn('w:sz'), str(size))
                left.set(qn('w:color'), color_hex)
                pBdr.append(left)
                pPr.append(pBdr)

            # Header (centered)
            name_paragraph = document.add_paragraph()
            format_paragraph(name_paragraph, line_spacing=1.1, space_after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            add_text(name_paragraph, resume_model.name, size=32, color=colors['text'])

            title_paragraph = document.add_paragraph()
            format_paragraph(title_paragraph, line_spacing=1.2, space_after=16, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            add_text(title_paragraph, resume_model.title, size=12, color=colors['muted'])

            contact_parts = [
                safe_text(resume_model.contact.phone),
                safe_text(resume_model.contact.email),
                safe_text(resume_model.contact.linkedin),
                safe_text(resume_model.contact.github)
            ]
            contact_parts = [item for item in contact_parts if item]
            if contact_parts:
                contact_paragraph = document.add_paragraph()
                format_paragraph(contact_paragraph, line_spacing=1.4, space_after=22, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                add_text(contact_paragraph, "   ".join(contact_parts), size=9, color=colors['muted'])

            def add_section_title(text):
                paragraph = document.add_paragraph()
                format_paragraph(paragraph, line_spacing=1.2, space_before=18, space_after=10)
                add_text(paragraph, text.upper(), size=11, color=colors['text'], bold=True)
                set_paragraph_bottom_border(paragraph, 'E0E0E0', size=4)
                return paragraph

            # Summary
            add_section_title("Summary")
            summary_paragraph = document.add_paragraph()
            format_paragraph(summary_paragraph, line_spacing=1.7, space_after=18, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
            add_text(summary_paragraph, resume_model.summary, size=10, color=colors['body'])

            # Skills
            add_section_title("Skills")
            skills_table = document.add_table(rows=0, cols=2)
            skills_table.autofit = False
            skills_table.columns[0].width = Inches(1.6)
            skills_table.columns[1].width = Inches(4.9)
            for skill in resume_model.skills:
                row = skills_table.add_row()
                left = row.cells[0]
                right = row.cells[1]
                self._remove_cell_borders(left)
                self._remove_cell_borders(right)
                left_paragraph = left.paragraphs[0]
                format_paragraph(left_paragraph, line_spacing=1.6)
                add_text(left_paragraph, safe_text(skill.category), size=11, color=colors['text'], bold=True)
                right_paragraph = right.paragraphs[0]
                format_paragraph(right_paragraph, line_spacing=1.6)
                add_text(right_paragraph, safe_text(skill.details), size=10, color=colors['body'])

            # Core Competencies
            add_section_title("Core Competencies")
            for competency in resume_model.core_competencies:
                title_paragraph = document.add_paragraph()
                format_paragraph(title_paragraph, line_spacing=1.4, space_after=2)
                add_text(title_paragraph, safe_text(competency.title), size=11, color=colors['text'], bold=True)
                desc_paragraph = document.add_paragraph()
                format_paragraph(desc_paragraph, line_spacing=1.6, space_after=10)
                add_text(desc_paragraph, safe_text(competency.description), size=9, color=colors['desc'])

            # Experience
            add_section_title("Experience")
            for job in resume_model.experience:
                role_paragraph = document.add_paragraph()
                format_paragraph(role_paragraph, line_spacing=1.3, space_after=2)
                add_text(role_paragraph, safe_text(job.role), size=12, color=colors['text'], bold=True)

                company_text = safe_text(job.company)
                if job.location:
                    company_text = f"{company_text} | {safe_text(job.location)}"
                company_paragraph = document.add_paragraph()
                format_paragraph(company_paragraph, line_spacing=1.3, space_after=2)
                add_text(company_paragraph, company_text, size=10, color=colors['body'])

                start = job.start_date.strftime('%B %Y') if job.start_date else ''
                end = job.end_date.strftime('%B %Y') if job.end_date else 'Present'
                meta_paragraph = document.add_paragraph()
                format_paragraph(meta_paragraph, line_spacing=1.3, space_after=6)
                add_text(meta_paragraph, f"{start} – {end}", size=9, color=colors['light'])

                if job.summary:
                    summary_paragraph = document.add_paragraph()
                    format_paragraph(summary_paragraph, line_spacing=1.7, space_after=8)
                    summary_paragraph.paragraph_format.left_indent = Inches(0.15)
                    summary_paragraph.paragraph_format.space_before = Pt(6)
                    set_paragraph_left_border(summary_paragraph, 'E0E0E0', size=6)
                    add_text(summary_paragraph, safe_text(job.summary), size=9, color=colors['desc'], italic=True)

                def add_subsection(title, items):
                    if not items:
                        return
                    heading = document.add_paragraph()
                    format_paragraph(heading, line_spacing=1.2, space_before=6, space_after=4)
                    add_text(heading, title.upper(), size=9, color=colors['muted'], bold=True)
                    for item in items:
                        bullet = document.add_paragraph(style='List Bullet')
                        format_paragraph(bullet, line_spacing=1.7, space_after=4)
                        add_text(bullet, item, size=9, color=colors['body'])

                if job.notable_projects:
                    add_subsection(
                        "Notable Projects",
                        [f"{safe_text(p.title)}: {safe_text(p.description)}" for p in job.notable_projects]
                    )

                add_subsection("Responsibilities", job.responsibilities or [])
                add_subsection("Accomplishments", job.accomplishments or [])

                if job.environment:
                    env_paragraph = document.add_paragraph()
                    format_paragraph(env_paragraph, line_spacing=1.6, space_after=12)
                    label_run = env_paragraph.add_run("Environment: ")
                    label_run.bold = True
                    label_run.font.size = Pt(9)
                    label_run.font.color.rgb = colors['text']
                    label_run.font.name = 'Calibri'
                    add_text(env_paragraph, join_values(job.environment), size=9, color=colors['muted'])

            # Education
            add_section_title("Education")
            for edu in resume_model.education:
                degree_paragraph = document.add_paragraph()
                format_paragraph(degree_paragraph, line_spacing=1.3, space_after=2)
                add_text(degree_paragraph, safe_text(edu.degree), size=11, color=colors['text'], bold=True)

                inst_text = safe_text(edu.institution)
                if edu.location:
                    inst_text = f"{inst_text}, {safe_text(edu.location)}"
                inst_paragraph = document.add_paragraph()
                format_paragraph(inst_paragraph, line_spacing=1.3, space_after=2)
                add_text(inst_paragraph, inst_text, size=10, color=colors['body'])

                meta_parts = []
                if edu.graduation_date:
                    meta_parts.append(edu.graduation_date.strftime('%B %Y'))
                if edu.gpa:
                    meta_parts.append(f"GPA: {safe_text(edu.gpa)}")
                if meta_parts:
                    meta_paragraph = document.add_paragraph()
                    format_paragraph(meta_paragraph, line_spacing=1.3, space_after=6)
                    add_text(meta_paragraph, " | ".join(meta_parts), size=9, color=colors['light'])

                if edu.relevance:
                    rel_paragraph = document.add_paragraph()
                    format_paragraph(rel_paragraph, line_spacing=1.6, space_after=10)
                    rel_paragraph.paragraph_format.left_indent = Inches(0.15)
                    set_paragraph_left_border(rel_paragraph, 'E0E0E0', size=6)
                    add_text(rel_paragraph, safe_text(edu.relevance), size=9, color=colors['desc'])

            # Certifications
            if resume_model.certifications:
                add_section_title("Certifications")
                for cert in resume_model.certifications:
                    name_paragraph = document.add_paragraph()
                    format_paragraph(name_paragraph, line_spacing=1.4, space_after=2)
                    add_text(name_paragraph, safe_text(cert.name), size=11, color=colors['text'], bold=True)
                    if cert.issue_date:
                        date_paragraph = document.add_paragraph()
                        format_paragraph(date_paragraph, line_spacing=1.4, space_after=2)
                        add_text(date_paragraph, cert.issue_date.strftime('%B %Y'), size=9, color=colors['light'])
                    if cert.description:
                        desc_paragraph = document.add_paragraph()
                        format_paragraph(desc_paragraph, line_spacing=1.6, space_after=8)
                        add_text(desc_paragraph, safe_text(cert.description), size=9, color=colors['desc'])

            docx_buffer = BytesIO()
            document.save(docx_buffer)
            docx_bytes = docx_buffer.getvalue()
            docx_buffer.close()

            if docx_bytes is None:
                raise Exception("DOCX generation failed")

            logger.info(f"Successfully generated minimal DOCX for {resume_model.name}")
            return docx_bytes

        except Exception as e:
            logger.error(f"Error generating minimal DOCX: {e}")
            raise Exception(f"Failed to generate minimal resume DOCX: {str(e)}")

    def _remove_leading_empty_paragraphs(self, document: Document):
        """Remove empty paragraphs at the beginning of the document and in table cells"""
        try:
            from docx.oxml import CT_P
            from docx.oxml.ns import qn

            body = document._element.body
            paragraphs_to_remove = []

            # Find leading empty paragraphs in document body
            for element in body:
                if element.tag == qn('w:p'):
                    # Check if paragraph is empty (no text content)
                    text = ''.join(node.text for node in element.iter() if hasattr(node, 'text') and node.text)
                    if not text or text.isspace():
                        paragraphs_to_remove.append(element)
                    else:
                        # Stop at first non-empty paragraph
                        break
                else:
                    # Stop at first non-paragraph element (like table)
                    break

            # Remove the empty paragraphs from body
            for para in paragraphs_to_remove:
                body.remove(para)

            if paragraphs_to_remove:
                logger.info(f"Removed {len(paragraphs_to_remove)} leading empty paragraph(s) from document body")

            # Also remove leading empty paragraphs from table cells
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_element = cell._element
                        cell_paras_to_remove = []

                        # Find leading empty paragraphs in this cell
                        for element in cell_element:
                            if element.tag == qn('w:p'):
                                text = ''.join(node.text for node in element.iter() if hasattr(node, 'text') and node.text)
                                # Check if it's empty and doesn't contain images
                                has_drawing = element.find(qn('w:r')) is not None and element.find('.//{}'.format(qn('w:drawing'))) is not None
                                if (not text or text.isspace()) and not has_drawing:
                                    cell_paras_to_remove.append(element)
                                else:
                                    # Stop at first non-empty paragraph or paragraph with image
                                    break

                        # Remove empty paragraphs from cell
                        for para in cell_paras_to_remove:
                            cell_element.remove(para)

        except Exception as e:
            logger.warning(f"Could not remove leading empty paragraphs: {e}")

    def _resize_logo_in_docx(self, document: Document):
        """Resize all images (logo) in the document to appropriate size"""
        try:
            from docx.shared import Inches

            target_width = Inches(0.7)

            # Search in both paragraphs and table cells
            def resize_images_in_element(element):
                for run in element.runs:
                    if hasattr(run, '_element'):
                        for drawing in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                            for inline in drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
                                extent = inline.find('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                                if extent is not None:
                                    cx = int(extent.get('cx'))
                                    cy = int(extent.get('cy'))
                                    aspect_ratio = cy / cx if cx > 0 else 1

                                    new_cx = int(target_width)
                                    new_cy = int(target_width * aspect_ratio)

                                    extent.set('cx', str(new_cx))
                                    extent.set('cy', str(new_cy))

                                    for ext in inline.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}ext'):
                                        ext.set('cx', str(new_cx))
                                        ext.set('cy', str(new_cy))

                                    logger.info(f"Resized logo to {target_width / 914400:.2f} inches")
                                    return True
                return False

            # Check paragraphs
            for paragraph in document.paragraphs:
                if resize_images_in_element(paragraph):
                    return

            # Check table cells
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if resize_images_in_element(paragraph):
                                return

        except Exception as e:
            logger.warning(f"Could not resize logo: {e}")

    def _fix_table_widths(self, document: Document):
        """Fix table column widths to match PDF template"""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from docx.shared import Inches

            # Standard page width for legal size with margins (8.5" - 2" margins = 6.5" usable)
            # In twips: 6.5 * 1440 = 9360
            page_width_twips = 9360

            skills_table_next = False

            for table_idx, table in enumerate(document.tables):
                # Get table text to identify which table this is
                table_text = table.rows[0].cells[0].text if table.rows else ""
                num_cols = len(table.rows[0].cells) if table.rows else 0
                num_rows = len(table.rows)

                if table_idx == 0:
                    # First table: Header table (logo + info)
                    # Logo: 0.7 inch = 1008 twips, increased to 1300 for better spacing
                    logo_width = 1300
                    info_width = page_width_twips - logo_width

                    self._set_table_column_widths(table, [logo_width, info_width])

                    # Remove borders from first column (logo cell)
                    if table.rows:
                        logo_cell = table.rows[0].cells[0]
                        self._remove_cell_borders(logo_cell)

                    logger.info(f"Set header table widths: logo={logo_width}, info={info_width}")

                elif "Key Expertise" in table_text or "Skills" in table_text:
                    # This is the skills title table
                    self._set_table_column_widths(table, [page_width_twips])
                    skills_table_next = True
                    logger.info("Found skills title table, next 2-column table will be skills data")

                elif skills_table_next and num_cols == 2 and num_rows > 1:
                    # This is the actual skills data table (right after the title)
                    # Skills table: 25% / 75% split
                    col1_width = int(page_width_twips * 0.25)
                    col2_width = page_width_twips - col1_width

                    self._set_table_column_widths(table, [col1_width, col2_width])
                    logger.info(f"Set skills table widths: category={col1_width}, details={col2_width}")
                    skills_table_next = False  # Reset flag

                elif num_cols == 1:
                    # Single column table (section title tables)
                    # Make them span full page width
                    self._set_table_column_widths(table, [page_width_twips])

                elif num_cols == 2 and num_rows == 2:
                    # Two row, two column tables (experience job entry headers)
                    col1_width = int(page_width_twips * 0.60)
                    col2_width = page_width_twips - col1_width
                    self._set_table_column_widths(table, [col1_width, col2_width])

        except Exception as e:
            logger.warning(f"Could not fix table widths: {e}")

    def _remove_cell_borders(self, cell):
        """Remove all borders from a table cell"""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            tcPr = cell._element.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                cell._element.insert(0, tcPr)

            # Remove existing borders
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is not None:
                tcPr.remove(tcBorders)

            # Add new borders with "nil" value (Word uses "nil" to hide borders)
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)

            tcPr.append(tcBorders)

        except Exception as e:
            logger.warning(f"Could not remove cell borders: {e}")

    def _set_table_column_widths(self, table, widths):
        """Set column widths for a table"""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            # Update table grid
            tbl = table._element
            tblGrid = tbl.find(qn('w:tblGrid'))

            if tblGrid is not None:
                # Remove old grid columns
                for gridCol in list(tblGrid):
                    tblGrid.remove(gridCol)

                # Add new grid columns with correct widths
                for width in widths:
                    gridCol = OxmlElement('w:gridCol')
                    gridCol.set(qn('w:w'), str(width))
                    tblGrid.append(gridCol)

            # Update cell widths in each row
            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    if idx < len(widths):
                        tcPr = cell._element.find(qn('w:tcPr'))
                        if tcPr is None:
                            tcPr = OxmlElement('w:tcPr')
                            cell._element.insert(0, tcPr)

                        tcW = tcPr.find(qn('w:tcW'))
                        if tcW is None:
                            tcW = OxmlElement('w:tcW')
                            tcPr.append(tcW)

                        tcW.set(qn('w:type'), 'dxa')
                        tcW.set(qn('w:w'), str(widths[idx]))

        except Exception as e:
            logger.warning(f"Could not set table column widths: {e}")

    def _apply_table_styling(self, document: Document, add_borders: bool = True):
        """Apply styling to tables for DOCX output"""
        try:
            if not add_borders:
                return

            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            for table in document.tables:
                # Apply table grid style for borders
                if table.style and 'Grid' not in str(table.style.name):
                    table.style = 'Table Grid'

                # Set table borders
                tbl = table._element
                tblPr = tbl.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = OxmlElement('w:tblPr')
                    tbl.insert(0, tblPr)

                # Add borders
                tblBorders = OxmlElement('w:tblBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:color'), '000000')
                    tblBorders.append(border)

                # Remove old borders and add new
                old_borders = tblPr.find(qn('w:tblBorders'))
                if old_borders is not None:
                    tblPr.remove(old_borders)
                tblPr.append(tblBorders)

        except Exception as e:
            logger.warning(f"Could not apply table styling: {e}")

    def _add_footer_to_docx(self, document: Document, footer_text: str):
        """Add footer to all pages of the document"""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from docx.shared import Pt

            section = document.sections[0]
            footer = section.footer

            # Clear existing footer content
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.clear()
            footer_para.alignment = 1  # Center alignment

            # Add top border to footer
            pPr = footer_para._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), '4')
            top.set(qn('w:space'), '1')
            top.set(qn('w:color'), '000000')
            pBdr.append(top)
            pPr.append(pBdr)

            # Add footer text
            formatted_text = footer_text.replace('Phone: ', '').replace(' | ', '   •   ')
            run = footer_para.add_run(formatted_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(8)

        except Exception as e:
            logger.warning(f"Could not add footer: {e}")

    def _apply_page_break_controls(self, document: Document):
        """Apply page break controls to prevent awkward breaks (similar to PDF template)"""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            # Apply to tables (prevent skills table from breaking)
            for table_idx, table in enumerate(document.tables):
                table_text = table.rows[0].cells[0].text if table.rows else ""

                # Skills table and section title tables should not break
                if "Key Expertise" in table_text or "Skills" in table_text:
                    # Find the actual skills data table (next table after title)
                    if len(table.rows) == 1 and len(table.rows[0].cells) == 1:
                        if table_idx + 1 < len(document.tables):
                            skills_table = document.tables[table_idx + 1]
                            # Prevent each row from splitting
                            for row in skills_table.rows:
                                self._set_row_cant_split(row)

                # Experience job entry tables - prevent rows from splitting
                elif len(table.rows) == 2 and len(table.rows[0].cells) == 2:
                    # This looks like a job entry header table
                    for row in table.rows:
                        self._set_row_cant_split(row)

            # Apply to paragraphs
            section_markers = [
                "Professional Summary",
                "Key Expertise",
                "Professional Experience",
                "Education",
                "Certifications"
            ]

            # Track if we're in Core Competencies section to avoid applying keep-with-next
            in_core_competencies = False

            for para_idx, para in enumerate(document.paragraphs):
                para_text = para.text.strip()

                # Check if entering or leaving Core Competencies section
                if "Core Competencies" in para_text:
                    in_core_competencies = True
                    self._set_keep_with_next(para)  # Keep title with content
                    continue
                elif any(marker in para_text for marker in section_markers):
                    in_core_competencies = False

                # Keep section titles with next paragraph
                if any(marker in para_text for marker in section_markers):
                    self._set_keep_with_next(para)

                # Keep company/institution names with next paragraph (but not in Core Competencies)
                if not in_core_competencies and para_idx + 1 < len(document.paragraphs):
                    next_para = document.paragraphs[para_idx + 1]
                    # If this is bold text followed by a heading-like paragraph, keep together
                    if para.runs and any(run.bold for run in para.runs):
                        self._set_keep_with_next(para)

            logger.info("Applied page break controls to document")

        except Exception as e:
            logger.warning(f"Could not apply page break controls: {e}")

    def _set_keep_with_next(self, paragraph):
        """Set paragraph to keep with next paragraph"""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            pPr = paragraph._element.get_or_add_pPr()
            keepNext = pPr.find(qn('w:keepNext'))
            if keepNext is None:
                keepNext = OxmlElement('w:keepNext')
                pPr.append(keepNext)
        except Exception as e:
            pass

    def _set_keep_lines_together(self, paragraph):
        """Set paragraph to keep all lines together"""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            pPr = paragraph._element.get_or_add_pPr()
            keepLines = pPr.find(qn('w:keepLines'))
            if keepLines is None:
                keepLines = OxmlElement('w:keepLines')
                pPr.append(keepLines)
        except Exception as e:
            pass

    def _set_row_cant_split(self, row):
        """Prevent table row from splitting across pages"""
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            trPr = row._element.find(qn('w:trPr'))
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                row._element.insert(0, trPr)

            cantSplit = trPr.find(qn('w:cantSplit'))
            if cantSplit is None:
                cantSplit = OxmlElement('w:cantSplit')
                trPr.append(cantSplit)
        except Exception as e:
            pass

    def _set_page_margins(self, document: Document):
        """Set page margins to Normal (1 inch all around) for legal size paper"""
        try:
            from docx.shared import Inches

            section = document.sections[0]

            # Set page size to legal (8.5" x 14")
            section.page_width = Inches(8.5)
            section.page_height = Inches(14)

            # Set margins to Normal (1 inch all around)
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

            logger.info("Set page margins to Normal (1 inch)")

        except Exception as e:
            logger.warning(f"Could not set page margins: {e}")

    def _set_default_font(self, document: Document, font_name: str):
        """Set default font for the entire document"""
        try:
            from docx.oxml.ns import qn

            # Set font for document defaults
            styles = document.styles
            style = styles['Normal']
            font = style.font
            font.name = font_name

            # Also set for East Asian and Complex scripts
            element = style.element
            rPr = element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                from docx.oxml import OxmlElement
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:cs'), font_name)

        except Exception as e:
            logger.warning(f"Could not set default font: {e}")
