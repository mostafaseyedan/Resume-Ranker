from google import genai
from google.genai import types
from pydantic import BaseModel
from typing import List
import io
import logging
import json

try:
    from docx import Document
    print("SUCCESS: python-docx imported successfully")
except ImportError as e:
    print(f"ERROR: Failed to import python-docx: {e}")
    Document = None

logger = logging.getLogger(__name__)

# Pydantic model for structured job extraction
class JobExtraction(BaseModel):
    job_title: str
    job_description_text: str
    required_skills: List[str]
    preferred_skills: List[str]
    experience_requirements: str
    education_requirements: List[str]
    certifications: List[str]
    key_responsibilities: List[str]
    soft_skills: List[str]
    other: List[str]

class GeminiFileProcessor:
    def __init__(self, api_key):
        self.client = genai.Client(api_key=api_key)
        self.supported_formats = ['.pdf', '.docx', '.doc']

    def upload_file_to_gemini(self, file):
        """Upload file directly to Gemini and return file object"""
        try:
            # Read file content
            file_content = file.read()
            if file_content is None:
                raise ValueError("Failed to read file content")
            file.seek(0)  # Reset file pointer

            # Create BytesIO object for upload
            file_data = io.BytesIO(file_content)

            # Determine MIME type based on file extension
            filename = file.filename.lower()
            if filename.endswith('.pdf'):
                mime_type = "application/pdf"
            elif filename.endswith(('.docx', '.doc')):
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            else:
                raise ValueError(f"Unsupported file format. Supported formats: {self.supported_formats}")

            # Upload file to Gemini
            uploaded_file = self.client.files.upload(
                file=file_data,
                config={"mime_type": mime_type}
            )

            logger.info(f"Successfully uploaded file {file.filename} to Gemini")
            return uploaded_file

        except Exception as e:
            logger.error(f"Error uploading file {file.filename} to Gemini: {e}")
            raise

    def extract_text_with_gemini(self, file):
        """Extract text from file - PDF via Gemini, DOCX via python-docx"""
        try:
            filename = file.filename.lower()

            if filename.endswith('.pdf'):
                # Use Gemini for PDF files
                uploaded_file = self.upload_file_to_gemini(file)

                prompt = "Extract all text content from this document. Preserve formatting and structure as much as possible. Return only the raw text content without any analysis or summary."

                response = self.client.models.generate_content(
                    model="models/gemini-flash-latest",
                    contents=[uploaded_file, prompt],
                    config=types.GenerateContentConfig(
                        thinking_config=types.ThinkingConfig(
                            thinking_budget=-1  # Dynamic thinking
                        )
                    )
                )

                extracted_text = response.text
                if extracted_text is None:
                    raise ValueError("Gemini response text is None")
                logger.info(f"Extracted {len(extracted_text)} characters from {file.filename} using Gemini")

            elif filename.endswith(('.docx', '.doc')):
                # Use python-docx for DOCX files
                if Document is None:
                    raise ValueError("python-docx library not available for DOCX processing")
                extracted_text = self._extract_from_docx(file)
                if extracted_text is None:
                    raise ValueError("DOCX extraction returned None")
                logger.info(f"Extracted {len(extracted_text)} characters from {file.filename} using python-docx")

            else:
                raise ValueError(f"Unsupported file format: {filename}")

            return extracted_text

        except Exception as e:
            logger.error(f"Error extracting text from {file.filename}: {e}")
            raise

    def _extract_from_docx(self, file):
        """Extract text from DOCX file using python-docx"""
        try:
            if Document is None:
                raise ValueError("python-docx library not available")

            # Read file content into memory
            file_content = file.read()
            if file_content is None:
                raise ValueError("Failed to read file content")
            file.seek(0)  # Reset file pointer

            # Open DOCX from memory
            document = Document(io.BytesIO(file_content))
            text = ""

            # Extract text from paragraphs
            for paragraph in document.paragraphs:
                text += paragraph.text + "\n"

            # Extract text from tables
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            return self._clean_text(text)

        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise

    def _clean_text(self, text):
        """Clean and normalize extracted text"""
        if not text:
            return ""

        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []

        for line in lines:
            line = line.strip()
            if line:  # Skip empty lines
                cleaned_lines.append(line)

        # Join lines with single newlines
        cleaned_text = '\n'.join(cleaned_lines)

        # Remove multiple spaces
        import re
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)

        return cleaned_text.strip()

    def extract_job_info_with_gemini(self, file):
        """Extract structured job information from PDF using Gemini Flash Lite"""
        try:
            # Upload file to Gemini
            uploaded_file = self.upload_file_to_gemini(file)

            # Extract structured job information using Gemini
            prompt = """
            Analyze this job description document and extract all relevant information including the job title, complete description text, required and preferred skills, experience requirements, education requirements, certifications, responsibilities, soft skills, and any other important details.
            """

            response = self.client.models.generate_content(
                model="models/gemini-flash-latest",
                contents=[uploaded_file, prompt],
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema=JobExtraction,
                    thinking_config=types.ThinkingConfig(
                        thinking_budget=-1  # Dynamic thinking
                    )
                )
            )

            if response.text is None or not response.text.strip():
                logger.error(f"Gemini returned empty response for {file.filename}")
                raise ValueError("Gemini response is empty")

            logger.info(f"Raw Gemini response for {file.filename}: {response.text[:200]}...")

            extracted_data = json.loads(response.text)
            logger.info(f"Extracted structured job info from {file.filename}: {extracted_data.get('job_title', 'Unknown Title')}")

            return extracted_data

        except Exception as e:
            logger.error(f"Error extracting job info from {file.filename} using Gemini: {e}")
            raise

    def validate_file(self, file):
        """Validate uploaded file"""
        if not file or file.filename == '':
            return False, "No file provided"

        if not any(file.filename.lower().endswith(ext) for ext in self.supported_formats):
            return False, f"Unsupported file format. Supported: {', '.join(self.supported_formats)}"

        # Check file size (Gemini supports up to 100MB, but we'll limit to 20MB for safety)
        file.seek(0, 2)  # Seek to end
        size = file.tell()
        file.seek(0)  # Reset to beginning

        if size > 20 * 1024 * 1024:  # 20MB limit
            return False, "File too large. Maximum size: 20MB"

        return True, "File is valid"
